"""Document-based ingestion source (PDF / DOCX).

This is the Phase 3 path: surveys not in Qualtrics (or that Rob doesn't own the
token for). We extract a single text payload and hand it to the normalizer; the
LLM does the heavy lifting of segmenting the document into question dicts.

We deliberately keep this dumb. Trying to parse PDF/DOCX layouts heuristically
is brittle; Claude handles messy text much better than regex.
"""
from __future__ import annotations

import io
from pathlib import Path


def extract_text(path: str | Path) -> str:
    path = Path(path)
    ext = path.suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(path)
    if ext in {".docx", ".doc"}:
        return _extract_docx(path)
    if ext in {".txt", ".md"}:
        return path.read_text(encoding="utf-8")
    raise ValueError(f"Unsupported document type for codebook ingest: {ext}")


def _extract_pdf(path: Path) -> str:
    try:
        import pdfplumber  # type: ignore
    except ImportError as e:
        raise ImportError(
            "pdfplumber is required for PDF ingestion. "
            "Run `pip install pdfplumber` and retry."
        ) from e

    chunks: list[str] = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            chunks.append(f"\n\n----- page {i} -----\n{text.strip()}")
    return "\n".join(chunks).strip()


def _extract_docx(path: Path) -> str:
    try:
        import docx  # type: ignore
    except ImportError as e:
        raise ImportError(
            "python-docx is required for DOCX ingestion. "
            "Run `pip install python-docx` and retry."
        ) from e

    doc = docx.Document(path)
    chunks = [p.text for p in doc.paragraphs if p.text.strip()]
    # Tables — common in survey docs for response scales
    for tbl in doc.tables:
        for row in tbl.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                chunks.append(row_text)
    return "\n".join(chunks).strip()
